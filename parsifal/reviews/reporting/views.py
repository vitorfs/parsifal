# coding: utf-8

import docx
from docx.shared import Inches

from django.contrib.auth.decorators import login_required
from django.core.urlresolvers import reverse as r
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse

from parsifal.reviews.models import *
from parsifal.reviews.decorators import author_required


@author_required
@login_required
def reporting(request, username, review_name):
    return redirect(r('export', args=(username, review_name,)))

@author_required
@login_required
def export(request, username, review_name):
    review = get_object_or_404(Review, name=review_name, author__username__iexact=username)
    return render(request, 'reporting/export.html', { 
            'review': review
        })

@author_required
@login_required
def download_docx(request):

    document = docx.Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph('first item in unordered list', style='ListBullet')
    document.add_paragraph('first item in ordered list', style='ListNumber')

    document.add_page_break()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    document.save(response)
    return response