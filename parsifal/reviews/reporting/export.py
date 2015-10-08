# coding: utf-8

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_review_to_docx(review, sections):
    document = Document()

    if 'name' in sections:
        h = document.add_heading(review.title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('')

    if 'authors' in sections:
        authors = list()
        authors.append(review.author.profile.get_screen_name())
        for author in review.co_authors.all():
            authors.append(author.profile.get_screen_name())
        p = document.add_paragraph(', '.join(authors))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('')

    if 'description' in sections:
        if review.description:
            document.add_paragraph(review.description)


    document.add_heading('Planning', level=2)

    if review.objective:
        document.add_paragraph(review.objective)

    '''
        PICOC
    '''
    if 'picoc' in sections:
        document.add_heading('PICOC', level=3)

        p = document.add_paragraph('', style='List Bullet')
        p.add_run('Population: ').bold = True
        p.add_run(review.population)

        p = document.add_paragraph('', style='List Bullet')
        p.add_run('Intervention: ').bold = True
        p.add_run(review.intervention)

        p = document.add_paragraph('', style='List Bullet')
        p.add_run('Comparison: ').bold = True
        p.add_run(review.comparison)

        p = document.add_paragraph('', style='List Bullet')
        p.add_run('Outcome: ').bold = True
        p.add_run(review.outcome)

        p = document.add_paragraph('', style='List Bullet')
        p.add_run('Context: ').bold = True
        p.add_run(review.context)

    '''
        Research Questions
    '''
    if 'research_questions' in sections:
        document.add_heading('Research Questions', level=3)

        for question in review.research_questions.all():
            document.add_paragraph(question.question, style='List Number')

    '''
        Keywords and Synonym
    '''
    if 'keywords_synonyms' in sections:
        document.add_heading('Keywords and Synonyms', level=3)

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Keyword'
        hdr_cells[1].text = 'Synonyms'

        for keyword in review.get_keywords():
            row_cells = table.add_row().cells
            row_cells[0].text = keyword.description
            row_cells[1].text = ', '.join(keyword.synonyms.all().values_list('description', flat=True))

    '''
        Search String
    '''
    if 'search_string' in sections:
        document.add_heading('Search String', level=3)
        document.add_paragraph(review.get_generic_search_string().search_string)

    '''
        Sources
    '''
    if 'sources' in sections:
        document.add_heading('Sources', level=3)

        for source in review.sources.all():
            text = source.name
            if source.url:
                text = u'{0} ({1})'.format(source.name, source.url)
            document.add_paragraph(text, style='List Bullet')

    '''
        Selection Criteria
    '''
    if 'selection_criteria' in sections:
        document.add_heading('Selection Criteria', level=3)

        p = document.add_paragraph()
        p.add_run('Inclusion Criteria:').bold = True
        for criteria in review.get_inclusion_criterias():
            document.add_paragraph(criteria.description, style='List Bullet')

        p = document.add_paragraph()
        p.add_run('Exclusion Criteria:').bold = True
        for criteria in review.get_exclusion_criterias():
            document.add_paragraph(criteria.description, style='List Bullet')

    '''
        Quality Assessment Checklist
    '''
    if 'quality_assessment_checklist' in sections:
        document.add_heading('Quality Assessment Checklist', level=3)

        p = document.add_paragraph()
        p.add_run('Questions:').bold = True
        for quality_question in review.get_quality_assessment_questions():
            document.add_paragraph(quality_question.description, style='List Bullet')

        p = document.add_paragraph()
        p.add_run('Answers:').bold = True
        for quality_answer in review.get_quality_assessment_answers():
            document.add_paragraph(quality_answer.description, style='List Bullet')

    '''
        Data Extraction Form
    '''
    if 'data_extraction_form' in sections:
        document.add_heading('Data Extraction Form', level=3)
        for field in review.get_data_extraction_fields():
            document.add_paragraph(field.description, style='List Bullet')

    '''
        Conducting
    '''

    document.add_heading('Conducting', level=2)

    '''
        Digital Libraries Search Strings
    '''

    if 'source_search_strings' in sections:
        document.add_heading('Digital Libraries Search Strings', level=3)
        for search_session in review.get_latest_source_search_strings():
            p = document.add_paragraph()
            p.add_run(u'{0}:'.format(search_session.source.name)).bold = True
            document.add_paragraph(search_session.search_string)
            document.add_paragraph()

    if 'number_imported_studies' in sections:
        document.add_heading('Imported Studies', level=3)
        for source in review.sources.all():
            p = document.add_paragraph(style='List Bullet')
            p.add_run(u'{0}: '.format(source.name)).bold = True
            count = review.article_set.filter(source=source).count()
            p.add_run(str(count))

    if 'quality_assessment' in sections:
        document.add_heading('Quality Assessment', level=3)

    if 'data_extraction' in sections:
        document.add_heading('Data Extraction', level=3)

    if 'data_analysis' in sections:
        document.add_heading('Data Analysis', level=3)

    return document
